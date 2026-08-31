from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "files" / "vladut-andrei-lambru-resume.docx"
INK = RGBColor(31, 35, 46)
MUTED = RGBColor(91, 98, 113)
ACCENT = RGBColor(91, 75, 180)
LINE = "D9DCE5"


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_font(run, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def add_link(paragraph, text, url, color=ACCENT):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    run.append(props)
    shade = OxmlElement("w:color")
    shade.set(qn("w:val"), str(color))
    props.append(shade)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rule(paragraph, color=LINE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def section_heading(container, title):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    set_repeat_font(run)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT
    add_rule(p)
    return p


def project(container, title, meta, sentences):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    set_repeat_font(r)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    r = p.add_run(f"  |  {meta}")
    set_repeat_font(r)
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    for sentence in sentences:
        p = container.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(.35)
        p.paragraph_format.first_line_indent = Cm(-.2)
        p.paragraph_format.space_after = Pt(.5)
        p.paragraph_format.line_spacing = 1.02
        r = p.add_run(sentence)
        set_repeat_font(r)
        r.font.size = Pt(8.7)
        r.font.color.rgb = INK


doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21)
section.top_margin = Cm(1.15)
section.bottom_margin = Cm(1.05)
section.left_margin = Cm(1.35)
section.right_margin = Cm(1.35)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(9)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(2)

name = doc.add_paragraph()
name.paragraph_format.space_after = Pt(0)
r = name.add_run("Vladut-Andrei Lambru")
set_repeat_font(r)
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = INK

role = doc.add_paragraph()
role.paragraph_format.space_after = Pt(4)
r = role.add_run("TECHNICAL GAME DESIGNER / GAMEPLAY PROGRAMMER")
set_repeat_font(r)
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = ACCENT
r.font.letter_spacing = Pt(.3)

contact = doc.add_paragraph()
contact.paragraph_format.space_after = Pt(7)
links = [
    ("v.lambru@st.hanze.nl", "mailto:v.lambru@st.hanze.nl"),
    ("vladut-andrei-lambru.github.io/Portfolio", "https://vladut-andrei-lambru.github.io/Portfolio/"),
    ("github.com/Vladut-Andrei-Lambru", "https://github.com/Vladut-Andrei-Lambru"),
    ("linkedin.com/in/vladut-andrei-lambru", "https://www.linkedin.com/in/vladut-andrei-lambru/"),
]
for index, (label, url) in enumerate(links):
    if index:
        r = contact.add_run("  ·  ")
        set_repeat_font(r)
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED
    add_link(contact, label, url)
add_rule(contact, color="9C8CFF", size="12")

profile = doc.add_paragraph()
profile.paragraph_format.space_after = Pt(4)
profile.paragraph_format.line_spacing = 1.08
r = profile.add_run(
    "CMGT student at Hanze University of Applied Sciences, currently completing an exchange at SeoulTech. "
    "I work mainly in Unity and C#, building gameplay, VR interaction, movement and feedback systems. "
    "Seeking a full-time internship from February 2027 in gameplay programming or technical game design."
)
set_repeat_font(r)
r.font.size = Pt(9.2)

section_heading(doc, "Selected projects")
project(doc, "Virtual Life Support", "Unity · C# · VR · 5-person team", [
    "Built hand-tracked chest compressions, ordered scenario logic, gestures, live rhythm/depth feedback and the results screen for a CPR practice scenario developed with Virtual Life Support.",
    "Turned interviews and tests with CPR-trained users into interactions; the prototype received 11/12 and was selected as the strongest of five student prototypes.",
])
project(doc, "Tiny Spider Tiny Home", "Unity · C# · Gameplay programming · 5-person team", [
    "Implemented movement across floors, walls and ceilings, a level third-person camera, web swinging, appliance interactions and a screen-space outline shader.",
    "Used surface normals and projected movement so one controller could handle every crawlable surface without separate movement modes.",
])
project(doc, "Maker's Fair", "Unity · C# · VR · Lead programmer", [
    "Built the runtime cart-construction mechanics, plank connections, hammer/nail interactions, guidance, UI and bridge-weight challenge.",
    "Iterated on unstable multi-part physics and added blueprints and placement holograms after playtests showed that the open building system needed clearer feedback.",
])

lower = doc.add_table(rows=1, cols=2)
lower.alignment = WD_TABLE_ALIGNMENT.CENTER
lower.autofit = False
lower.columns[0].width = Cm(10.8)
lower.columns[1].width = Cm(7.2)
left, right = lower.rows[0].cells
for cell in (left, right):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=40, bottom=0, start=0, end=180 if cell is left else 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tc_pr.append(borders)

section_heading(left, "Education")
for title, place, date in [
    ("Creative Media & Game Technologies, BSc", "Hanze University of Applied Sciences · Groningen", "2024-2028"),
    ("Computer Science exchange", "SeoulTech · Seoul", "Aug 2026-Jan 2027"),
]:
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    set_repeat_font(r); r.bold = True; r.font.size = Pt(8.8)
    r = p.add_run(f"  {date}")
    set_repeat_font(r); r.font.size = Pt(8); r.font.color.rgb = MUTED
    p = left.add_paragraph(place)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: set_repeat_font(r); r.font.size = Pt(8.2); r.font.color.rgb = MUTED

section_heading(left, "Additional experience")
p = left.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("Sales employee, Albert Heijn")
set_repeat_font(r); r.bold = True; r.font.size = Pt(8.7)
r = p.add_run(" · Groningen · 2024-present")
set_repeat_font(r); r.font.size = Pt(8); r.font.color.rgb = MUTED
p = left.add_paragraph("Part-time work alongside university: customer support, restocking and deliveries.")
for r in p.runs: set_repeat_font(r); r.font.size = Pt(8.2)
p = left.add_paragraph("Festival volunteer · Neversea 2023 · Beach, Please! 2024")
for r in p.runs: set_repeat_font(r); r.font.size = Pt(8.2); r.font.color.rgb = MUTED

section_heading(right, "Technical skills")
skills = [
    ("Core", "Unity, C#, gameplay programming, technical design, Git"),
    ("Systems", "VR/XR interaction, physics, character controllers, cameras, UI/UX"),
    ("Also used", "Unreal Engine, Blueprints, C++, GameMaker"),
]
for label, value in skills:
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    set_repeat_font(r); r.bold = True; r.font.size = Pt(8.4)
    r = p.add_run(value)
    set_repeat_font(r); r.font.size = Pt(8.4)

section_heading(right, "Certificates")
for line in [
    "Propedeutic Diploma · Hanze · 2025",
    "Cambridge English · score 173 · 2024",
    "Oracle Academy Database Design · 2023",
]:
    p = right.add_paragraph(line)
    p.paragraph_format.space_after = Pt(1)
    for r in p.runs: set_repeat_font(r); r.font.size = Pt(8.2)

section_heading(right, "Languages")
p = right.add_paragraph("Romanian · Native\nEnglish · Professional working proficiency")
for r in p.runs: set_repeat_font(r); r.font.size = Pt(8.2)

doc.core_properties.title = "Vladut-Andrei Lambru - Technical Game Designer and Gameplay Programmer"
doc.core_properties.subject = "Internship CV"
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.save(OUT)
print(OUT)
